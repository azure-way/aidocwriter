from __future__ import annotations

import json
import re
import uuid
from dataclasses import dataclass
import logging
import os
from pathlib import Path
from typing import Any, Dict, Optional, Tuple, List
from concurrent.futures import ThreadPoolExecutor
import time

try:  # Optional dependency for diagram rendering
    import requests  # type: ignore
except Exception:  # pragma: no cover
    requests = None  # type: ignore

try:  # optional dependency until queue features are used
    from azure.servicebus import AutoLockRenewer, ServiceBusClient, ServiceBusMessage  # type: ignore
except Exception:  # pragma: no cover
    ServiceBusClient = None  # type: ignore
    ServiceBusMessage = None  # type: ignore
    AutoLockRenewer = None  # type: ignore

from .config import get_settings
from .agents.planner import PlannerAgent
from .agents.writer import WriterAgent
from .agents.reviewer import ReviewerAgent
from .agents.verifier import VerifierAgent
from .summary import Summarizer
from .graph import build_dependency_graph
from .storage import BlobStore
from .telemetry import stage_timer
from .agents.interviewer import InterviewerAgent
from .agents.style_reviewer import StyleReviewerAgent
from .agents.cohesion_reviewer import CohesionReviewerAgent
from .agents.summary_reviewer import SummaryReviewerAgent


LOG_CONFIGURED = False


def _configure_logging(worker_name: str) -> None:
    global LOG_CONFIGURED
    if LOG_CONFIGURED:
        return
    log_level = os.getenv("DOCWRITER_LOG_LEVEL", "INFO").upper()
    formatter = logging.Formatter("%(asctime)s %(levelname)s [%(name)s] %(message)s")
    handlers: list[logging.Handler] = []

    log_dir = os.getenv("LOG_DIR")
    if log_dir:
        log_path = Path(log_dir)
        log_path.mkdir(parents=True, exist_ok=True)
        file_handler = logging.FileHandler(log_path / f"{worker_name}.log")
        file_handler.setFormatter(formatter)
        handlers.append(file_handler)

    stream_handler = logging.StreamHandler()
    stream_handler.setFormatter(formatter)
    handlers.append(stream_handler)

    logging.basicConfig(level=log_level, handlers=handlers, force=True)
    LOG_CONFIGURED = True


@dataclass
class Job:
    title: str
    audience: str
    out: str = ""
    job_id: str | None = None
    cycles: int = 1


def _sb_check():
    settings = get_settings()
    if not settings.sb_connection_string:
        raise RuntimeError("SERVICE_BUS_CONNECTION_STRING not set")
    if ServiceBusClient is None or ServiceBusMessage is None:
        raise RuntimeError("azure-servicebus not installed. Install with `pip install azure-servicebus`.\n")
    return settings


def _send(queue_name: str, payload: Dict[str, Any]) -> None:
    settings = _sb_check()
    with ServiceBusClient.from_connection_string(settings.sb_connection_string) as client:
        with client.get_queue_sender(queue_name) as sender:
            sender.send_messages(ServiceBusMessage(json.dumps(payload)))


def _status(payload: Dict[str, Any]) -> None:
    settings = _sb_check()
    try:
        with ServiceBusClient.from_connection_string(settings.sb_connection_string) as client:
            with client.get_topic_sender(settings.sb_topic_status) as sender:
                sender.send_messages(ServiceBusMessage(json.dumps(payload)))
    except Exception:
        pass
    try:  # mirror to API status store if available
        from api.status_store import status_store  # type: ignore

        status_store.record(payload)
    except Exception:
        pass


def send_job(job: Job) -> str:
    job_id = job.job_id or str(uuid.uuid4())
    try:
        store = BlobStore()
        blob_path = store.allocate_document_blob(job_id)
    except Exception:
        blob_path = job.out or f"/tmp/{job_id}_document.md"
    payload = {
        "job_id": job_id,
        "title": job.title,
        "audience": job.audience,
        "out": blob_path,
        "cycles_remaining": max(1, int(job.cycles)),
        "cycles_completed": 0,
    }
    settings = get_settings()
    _send(settings.sb_queue_plan_intake, payload)
    _status({"job_id": job_id, "stage": "ENQUEUED", "ts": time.time()})
    return job_id


def worker_plan_intake() -> None:
    settings = _sb_check()
    interviewer = InterviewerAgent()

    def handle(_msg, data: Dict[str, Any]):
        process_plan_intake(data, interviewer)
    _run_processor(settings.sb_queue_plan_intake, handle)


def worker_intake_resume() -> None:
    settings = _sb_check()

    def handle(_msg, data: Dict[str, Any]):
        process_intake_resume(data)

    _run_processor(settings.sb_queue_intake_resume, handle)


def send_resume(job_id: str) -> None:
    settings = get_settings()
    payload: Dict[str, Any] = {"job_id": job_id}
    try:
        store = BlobStore()
        context_text = store.get_text(blob=f"jobs/{job_id}/intake/context.json")
        context = json.loads(context_text)
        if isinstance(context, dict):
            payload.update(
                {
                    "title": context.get("title"),
                    "audience": context.get("audience"),
                    "out": context.get("out"),
                    "cycles_remaining": context.get("cycles_remaining"),
                    "cycles_completed": context.get("cycles_completed"),
                }
            )
    except Exception:
        pass
    if not isinstance(payload.get("out"), str) or not payload.get("out"):
        try:
            payload["out"] = BlobStore().allocate_document_blob(job_id)
        except Exception:
            payload["out"] = f"/tmp/{job_id}_document.md"
    _send(settings.sb_queue_intake_resume, payload)


def _run_processor(queue_name: str, handler, max_workers: int = 1) -> None:
    settings = _sb_check()
    renew_seconds = getattr(settings, "sb_lock_renew_s", 0)
    lock_renewer = None
    if AutoLockRenewer is not None and renew_seconds and renew_seconds > 0:
        lock_renewer = AutoLockRenewer(max_lock_renewal_duration=renew_seconds)
    try:
        with ServiceBusClient.from_connection_string(settings.sb_connection_string) as client:
            with client.get_queue_receiver(queue_name, max_wait_time=30) as receiver:
                with ThreadPoolExecutor(max_workers=max_workers) as pool:
                    while True:
                        try:
                            messages = receiver.receive_messages(max_message_count=10, max_wait_time=30)
                        except Exception:
                            time.sleep(2)
                            continue
                        if not messages:
                            continue
                        for msg in messages:
                            try:
                                data = _decode_msg(msg)
                            except Exception:
                                receiver.abandon_message(msg)
                                continue
                            if lock_renewer is not None:
                                try:
                                    lock_renewer.register(
                                        receiver,
                                        msg,
                                        max_lock_renewal_duration=renew_seconds,
                                    )
                                except Exception as renew_exc:
                                    logging.exception(
                                        "Failed to register auto lock renewal for message %s: %s",
                                        msg.message_id,
                                        renew_exc,
                                    )
                            fut = pool.submit(handler, msg, data)
                            try:
                                fut.result()
                                receiver.complete_message(msg)
                            except Exception as ex:
                                logging.exception(
                                    f"Message processing failed, abandoning message {msg.message_id}: {ex}"
                                )
                                receiver.abandon_message(msg)
    finally:
        if lock_renewer is not None:
            lock_renewer.close()


def worker_plan() -> None:
    _configure_logging("worker-plan")
    settings = _sb_check()
    planner = PlannerAgent()

    def handle(_msg, data: Dict[str, Any]):
        process_plan(data, planner)

    _run_processor(settings.sb_queue_plan, handle)


def worker_write() -> None:
    _configure_logging("worker-write")
    settings = _sb_check()
    writer = WriterAgent()
    summarizer = Summarizer()

    def handle(_msg, data: Dict[str, Any]):
        process_write(data, writer, summarizer)

    _run_processor(settings.sb_queue_write, handle)


def worker_review() -> None:
    _configure_logging("worker-review")
    settings = _sb_check()
    reviewer = ReviewerAgent()

    def handle(_msg, data: Dict[str, Any]):
        process_review(data, reviewer)

    _run_processor(settings.sb_queue_review, handle)


def worker_verify() -> None:
    _configure_logging("worker-verify")
    settings = _sb_check()
    verifier = VerifierAgent()

    def handle(_msg, data: Dict[str, Any]):
        process_verify(data, verifier)

    _run_processor(settings.sb_queue_verify, handle)


def worker_rewrite() -> None:
    _configure_logging("worker-rewrite")
    settings = _sb_check()
    writer = WriterAgent()

    def handle(_msg, data: Dict[str, Any]):
        process_rewrite(data, writer)

    _run_processor(settings.sb_queue_rewrite, handle)


def worker_finalize() -> None:
    _configure_logging("worker-finalize")
    settings = _sb_check()

    def handle(_msg, data: Dict[str, Any]):
        process_finalize(data)

    _run_processor(settings.sb_queue_finalize, handle)


def _decode_msg(msg) -> Dict[str, Any]:
    try:
        return json.loads(str(msg))
    except Exception:
        return json.loads("".join([b.decode("utf-8") for b in msg.body]))


# Exposed per-stage processors for E2E tests and direct invocation
def process_plan_intake(data: Dict[str, Any], interviewer: InterviewerAgent | None = None) -> None:
    interviewer = interviewer or InterviewerAgent()
    with stage_timer(job_id=data["job_id"], stage="PLAN_INTAKE"):
        title = data["title"]
        questions = interviewer.propose_questions(title)
        try:
            store = BlobStore()
            store.put_text(
                blob=f"jobs/{data['job_id']}/intake/questions.json", text=json.dumps(questions, indent=2)
            )
            store.put_text(
                blob=f"jobs/{data['job_id']}/intake/context.json", text=json.dumps(data, indent=2)
            )
        except Exception:
            pass
    _status(
        {
            "job_id": data["job_id"],
            "stage": "INTAKE_READY",
            "artifact": f"jobs/{data['job_id']}/intake/questions.json",
            "message": "Upload answers.json to intake folder and call resume",
            "ts": time.time(),
        }
    )
    # Do not advance the pipeline until answers are provided via resume.


def process_intake_resume(data: Dict[str, Any]) -> None:
    settings = get_settings()
    with stage_timer(job_id=data["job_id"], stage="INTAKE_RESUME"):
        _send(settings.sb_queue_plan, data)
        _status({"job_id": data["job_id"], "stage": "INTAKE_RESUMED", "ts": time.time()})


def process_plan(data: Dict[str, Any], planner: PlannerAgent | None = None) -> None:
    settings = get_settings()
    planner = planner or PlannerAgent()
    store: BlobStore | None = None
    try:
        store = BlobStore()
    except Exception:
        store = None
    print(
        "[worker-plan] Processing job",
        data.get("job_id"),
        "title=",
        data.get("title"),
    )
    with stage_timer(job_id=data["job_id"], stage="PLAN"):
        audience = data.get("audience")
        title = data.get("title")
        length_pages = 80
        
        answers: Dict[str, Any] = {}
        if store:
            try:
                plan_text = store.get_text(blob=f"jobs/{data['job_id']}/plan.json")
                existing_plan = json.loads(plan_text)
                title = existing_plan.get("title", title)
                audience = existing_plan.get("audience", audience)
                if existing_plan.get("length_pages") is not None:
                    length_pages = int(existing_plan.get("length_pages"))
            except Exception:
                pass

            try:
                answers_text = store.get_text(blob=f"jobs/{data['job_id']}/intake/answers.json")
                answers = json.loads(answers_text)
                audience = answers.get("audience", audience)
                title = answers.get("title", title)
                length_pages = int(answers.get("length_pages", length_pages))
            except Exception:
                pass

        plan = planner.plan(title or "", audience=audience or "", length_pages=length_pages)
        try:
            plan.global_style.update({
                "tone": answers.get("tone") or plan.global_style.get("tone"),
                "pov": answers.get("pov") or plan.global_style.get("pov"),
                "structure": answers.get("structure") or plan.global_style.get("structure"),
                "constraints": answers.get("constraints") or plan.global_style.get("constraints"),
            })
        except Exception:
            pass
    payload = {
        **data,
        "plan": {
            "title": plan.title,
            "audience": plan.audience,
            "length_pages": max(60, plan.length_pages or 80),
            "outline": plan.outline,
            "glossary": plan.glossary,
            "global_style": plan.global_style,
            "diagram_specs": plan.diagram_specs,
        },
        "dependency_summaries": {},
        "intake_answers": answers,
    }
    try:
        target_store = store or BlobStore()
        target_store.put_text(
            blob=f"jobs/{data['job_id']}/plan.json",
            text=json.dumps(payload["plan"], indent=2),
        )
    except Exception:
        pass
    _send(settings.sb_queue_write, payload)
    _status({"job_id": data["job_id"], "stage": "PLAN_DONE", "artifact": f"jobs/{data['job_id']}/plan.json", "ts": time.time()})
    print("[worker-plan] Dispatched job", data.get("job_id"), "to writing queue")


def process_write(data: Dict[str, Any], writer: WriterAgent | None = None, summarizer: Summarizer | None = None) -> None:
    settings = get_settings()
    writer = writer or WriterAgent()
    summarizer = summarizer or Summarizer()
    blob_path = data.get("out")
    if not isinstance(blob_path, str) or not blob_path:
        blob_path = BlobStore().allocate_document_blob(data["job_id"])

    with stage_timer(job_id=data["job_id"], stage="WRITE"):
        plan = data["plan"]
        outline = plan.get("outline", [])
        graph = build_dependency_graph(outline)
        order = graph.topological_order() if outline else []
        id_to_section = {str(s.get("id")): s for s in outline}
        dependency_summaries = data.get("dependency_summaries", {})
        document_text_parts: list[str] = []
        for sid in order:
            section = id_to_section[sid]
            deps = section.get("dependencies", []) or []
            dep_context = "\n".join([dependency_summaries.get(str(d), "") for d in deps if dependency_summaries.get(str(d))])
            section_output = "".join(list(writer.write_section(plan=plan, section=section, dependency_context=dep_context)))
            document_text_parts.append(section_output)
            summary = summarizer.summarize_section("\n\n".join(document_text_parts))
            dependency_summaries[sid] = summary
        document_text = "\n\n".join(document_text_parts)
    payload = {**data, "out": blob_path, "dependency_summaries": dependency_summaries}
    try:
        store = BlobStore()
        store.put_text(blob=blob_path, text=document_text)
        store.put_text(blob=f"jobs/{data['job_id']}/draft.md", text=document_text)
    except Exception:
        pass
    _send(settings.sb_queue_review, payload)
    _status({"job_id": data["job_id"], "stage": "WRITE_DONE", "artifact": f"jobs/{data['job_id']}/draft.md", "ts": time.time()})


def process_review(data: Dict[str, Any], reviewer: ReviewerAgent | None = None) -> None:
    settings = get_settings()
    reviewer = reviewer or ReviewerAgent()
    with stage_timer(job_id=data["job_id"], stage="REVIEW", cycle=int(data.get("cycles_completed", 0)) + 1):
        store = BlobStore()
        draft = store.get_text(blob=data["out"])
        review_json = reviewer.review(plan=data["plan"], draft_markdown=draft)
        style = StyleReviewerAgent().review_style(plan=data["plan"], markdown=draft)
        cohesion = CohesionReviewerAgent().review_cohesion(plan=data["plan"], markdown=draft)
        summary = SummaryReviewerAgent().review_executive_summary(plan=data["plan"], markdown=draft)
        cycle_idx = int(data.get("cycles_completed", 0)) + 1
    payload = {**data, "review_json": review_json, "style_json": style, "cohesion_json": cohesion, "exec_summary_json": summary}
    try:
        store = BlobStore()
        store.put_text(
            blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/review.json",
            text=review_json,
        )
        store.put_text(blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/style.json", text=style)
        store.put_text(blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/cohesion.json", text=cohesion)
        store.put_text(blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/executive_summary.json", text=summary)
    except Exception:
        pass
    _send(settings.sb_queue_verify, payload)
    _status({"job_id": data["job_id"], "stage": "REVIEW_DONE", "ts": time.time(), "cycle": cycle_idx})


def process_verify(data: Dict[str, Any], verifier: VerifierAgent | None = None) -> None:
    settings = get_settings()
    verifier = verifier or VerifierAgent()
    with stage_timer(job_id=data["job_id"], stage="VERIFY", cycle=int(data.get("cycles_completed", 0)) + 1):
        store = BlobStore()
        draft = store.get_text(blob=data["out"])
        cycle_idx = int(data.get("cycles_completed", 0)) + 1
        try:
            review_data = json.loads(data.get("review_json", "{}"))
            revised = review_data.get("revised_markdown")
            if isinstance(revised, str) and revised.strip():
                merged = _merge_revised_markdown(draft, revised)
                if merged != draft:
                    draft = merged
                    try:
                        store.put_text(blob=data["out"], text=merged)
                        store.put_text(
                            blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/revision.md",
                            text=merged,
                        )
                    except Exception:
                        pass
        except Exception:
            pass
        placeholder_sections = _find_placeholder_sections(draft)
        verification_json = verifier.verify(
            dependency_summaries=data.get("dependency_summaries", {}), final_markdown=draft
        )
    payload = {**data, "verification_json": verification_json}
    try:
        store = BlobStore()
        store.put_text(
            blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/contradictions.json",
            text=verification_json,
        )
    except Exception:
        pass
    try:
        verification = json.loads(verification_json)
        contradictions = verification.get("contradictions", [])
    except Exception:
        contradictions = []

    style_guidance, style_sections = _parse_review_guidance(data.get("style_json"))
    cohesion_guidance, cohesion_sections = _parse_review_guidance(data.get("cohesion_json"))
    needs_rewrite = (
        bool(contradictions)
        or bool(style_guidance)
        or bool(cohesion_guidance)
        or bool(placeholder_sections)
    )

    if needs_rewrite and int(payload.get("cycles_remaining", 0)) > 0:
        payload["placeholder_sections"] = sorted(placeholder_sections)
        _send(settings.sb_queue_rewrite, payload)
    else:
        _send(settings.sb_queue_finalize, payload)
    _status(
        {
            "job_id": data["job_id"],
            "stage": "VERIFY_DONE",
            "has_contradictions": bool(contradictions),
            "style_issues": bool(style_guidance),
            "cohesion_issues": bool(cohesion_guidance),
            "placeholder_sections": bool(placeholder_sections),
            "cycle": cycle_idx,
            "ts": time.time(),
        }
    )


def process_rewrite(data: Dict[str, Any], writer: WriterAgent | None = None) -> None:
    settings = get_settings()
    writer = writer or WriterAgent()
    with stage_timer(job_id=data["job_id"], stage="REWRITE", cycle=int(data.get("cycles_completed", 0)) + 1):
        plan = data["plan"]
        store = BlobStore()
        text = store.get_text(blob=data["out"])
        cycle_idx = int(data.get("cycles_completed", 0)) + 1
        try:
            verification = json.loads(data.get("verification_json", "{}"))
        except Exception:
            verification = {"contradictions": []}
        contradictions = verification.get("contradictions", [])
        id_to_section = {str(s.get("id")): s for s in plan.get("outline", [])}
        dependency_summaries = data.get("dependency_summaries", {})

        style_guidance, style_sections = _parse_review_guidance(data.get("style_json"))
        cohesion_guidance, cohesion_sections = _parse_review_guidance(data.get("cohesion_json"))
        combined_guidance = "\n".join(filter(None, [style_guidance, cohesion_guidance]))

        affected = {str(c.get("section_id")) for c in contradictions if c.get("section_id")}
        affected.update(style_sections)
        affected.update(cohesion_sections)

        if not affected and combined_guidance:
            affected = set(id_to_section.keys())

        placeholder_sections = {str(s) for s in data.get("placeholder_sections", [])}
        affected.update(placeholder_sections)

        if affected:
            for sid in affected:
                section = id_to_section.get(sid)
                if not section:
                    continue
                deps = section.get("dependencies", []) or []
                dep_context = "\n".join(
                    [dependency_summaries.get(str(d), "") for d in deps if dependency_summaries.get(str(d))]
                )
                new_text = "".join(
                    list(
                        writer.write_section(
                            plan=plan,
                            section=section,
                            dependency_context=dep_context,
                            extra_guidance=combined_guidance,
                        )
                    )
                )
                start_marker = f"<!-- SECTION:{sid}:START -->"
                end_marker = f"<!-- SECTION:{sid}:END -->"
                start_idx = text.find(start_marker)
                end_idx = text.find(end_marker)
                if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                    end_idx += len(end_marker)
                    text = text[:start_idx] + new_text + text[end_idx:]
            store.put_text(blob=data["out"], text=text)
            try:
                store.put_text(blob=f"jobs/{data['job_id']}/cycle_{cycle_idx}/rewrite.md", text=text)
            except Exception:
                pass
    payload = {
        **data,
        "cycles_remaining": max(0, int(data.get("cycles_remaining", 0)) - 1),
        "cycles_completed": int(data.get("cycles_completed", 0)) + 1,
        "placeholder_sections": [],
    }
    _send(settings.sb_queue_review, payload)
    _status({"job_id": data["job_id"], "stage": "REWRITE_DONE", "cycle": cycle_idx, "ts": time.time()})


def process_finalize(data: Dict[str, Any]) -> None:
    with stage_timer(job_id=data["job_id"], stage="FINALIZE"):
        try:
            store = BlobStore()
            final_text = store.get_text(blob=data["out"])
            tmp_dir = Path(tempfile.mkdtemp())
            markdown_path = tmp_dir / "document.md"
            markdown_path.write_text(final_text, encoding="utf-8")
            final_text, image_paths = _replace_mermaid_with_images(final_text, markdown_path)
            markdown_path.write_text(final_text, encoding="utf-8")
            store.put_text(blob=data["out"], text=final_text)
            store.put_text(blob=f"jobs/{data['job_id']}/final.md", text=final_text)
            for idx, img_path in enumerate(image_paths, start=1):
                try:
                    store.put_bytes(
                        blob=f"jobs/{data['job_id']}/images/diagram_{idx}{img_path.suffix}",
                        data_bytes=img_path.read_bytes(),
                    )
                except Exception:
                    logging.exception("Failed to upload diagram %s for job %s", img_path, data.get("job_id"))
            exports: Dict[str, Path] = {}
            pdf_path = _export_pdf(markdown_path)
            if pdf_path:
                exports["pdf"] = pdf_path
            docx_path = _export_docx(markdown_path)
            if docx_path:
                exports["docx"] = docx_path
            for fmt, path in exports.items():
                try:
                    store.put_bytes(
                        blob=f"jobs/{data['job_id']}/final.{fmt}",
                        data_bytes=path.read_bytes(),
                    )
                except Exception:
                    logging.exception("Failed to upload %s export for job %s", fmt, data.get("job_id"))
        except Exception:
            pass
    _status({"job_id": data["job_id"], "stage": "FINALIZE_DONE", "ts": time.time()})
SECTION_START_RE = re.compile(r"<!-- SECTION:(?P<id>[^:]+):START -->")


def _extract_sections(text: str) -> Dict[str, str]:
    sections: Dict[str, str] = {}
    for match in SECTION_START_RE.finditer(text):
        sid = match.group("id")
        start_idx = match.start()
        end_marker = f"<!-- SECTION:{sid}:END -->"
        end_idx = text.find(end_marker, match.end())
        if end_idx == -1:
            continue
        end_idx += len(end_marker)
        sections[sid] = text[start_idx:end_idx]
    return sections


def _merge_revised_markdown(original: str, revised: str) -> str:
    if not revised.strip():
        return original
    revised_sections = _extract_sections(revised)
    if not revised_sections:
        # Assume full-document replacement
        return revised
    original_sections = _extract_sections(original)
    if not original_sections:
        return revised
    updated = original
    for sid, section_text in revised_sections.items():
        original_section = original_sections.get(sid)
        if not original_section:
            continue
        inner = section_text.replace(f"<!-- SECTION:{sid}:START -->", "").replace(
            f"<!-- SECTION:{sid}:END -->", ""
        ).strip()
        if not inner or "content unchanged" in inner.lower():
            continue
        updated = updated.replace(original_section, section_text)
    return updated


def _parse_review_guidance(raw: Any) -> Tuple[str, set[str]]:
    if not isinstance(raw, str):
        return "", set()
    raw = raw.strip()
    if not raw:
        return "", set()
    sections: set[str] = set()
    try:
        parsed = json.loads(raw)
    except Exception:
        return raw, sections
    lines: list[str] = []

    def _handle_item(label: str, value: Any) -> None:
        if isinstance(value, dict):
            section_id = value.get("section_id") or value.get("section") or value.get("id")
            if section_id:
                sections.add(str(section_id))
            description = value.get("description") or value.get("issue") or value.get("summary")
            if not description:
                description = json.dumps(value, ensure_ascii=False)
            lines.append(f"{label}: {description}")
        else:
            lines.append(f"{label}: {value}")

    if isinstance(parsed, dict):
        for key, val in parsed.items():
            if isinstance(val, list):
                for item in val:
                    _handle_item(key, item)
            else:
                _handle_item(key, val)
    elif isinstance(parsed, list):
        for item in parsed:
            _handle_item("item", item)
    else:
        lines.append(str(parsed))

    guidance_text = "\n".join(line for line in lines if line).strip()
    if not guidance_text:
        guidance_text = json.dumps(parsed, ensure_ascii=False)
    return guidance_text, sections


def _find_placeholder_sections(markdown: str) -> set[str]:
    placeholders: set[str] = set()
    sections = _extract_sections(markdown)
    for sid, section_text in sections.items():
        inner = section_text.replace(f"<!-- SECTION:{sid}:START -->", "").replace(
            f"<!-- SECTION:{sid}:END -->", ""
        ).strip()
        inner_lower = inner.lower()
        if "content unchanged" in inner_lower or "placeholder" in inner_lower:
            placeholders.add(sid)
    return placeholders


def _markdown_to_html(markdown: str) -> str:
    try:
        from markdown_it import MarkdownIt
    except Exception:
        logging.warning("markdown-it-py not installed; skipping HTML conversion")
        return ""

    md = MarkdownIt("commonmark", {"html": True, "linkify": True})
    md.enable("table")
    md.enable("strikethrough")
    return md.render(markdown)


def _export_pdf(markdown_path: Path) -> Optional[Path]:
    try:
        html = _markdown_to_html(markdown_path.read_text(encoding="utf-8"))
        if not html:
            return None
        from weasyprint import HTML  # type: ignore
    except Exception:
        logging.warning("WeasyPrint or markdown-it-py not available; skipping PDF export")
        return None

    pdf_path = markdown_path.with_suffix(".pdf")
    try:
        HTML(string=html, base_url=str(markdown_path.parent)).write_pdf(str(pdf_path))
        return pdf_path
    except Exception:
        logging.exception("Failed to render PDF for %s", markdown_path)
        return None


IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")


def _export_docx(markdown_path: Path) -> Optional[Path]:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception:
        logging.warning("python-docx not installed; skipping DOCX export")
        return None

    try:
        lines = markdown_path.read_text(encoding="utf-8").splitlines()
    except Exception:
        logging.exception("Failed to read markdown for DOCX export: %s", markdown_path)
        return None

    doc = Document()

    def _add_image(image_src: str) -> None:
        img_path = (markdown_path.parent / image_src).resolve()
        if not img_path.exists():
            logging.warning("Image %s not found for DOCX export", img_path)
            return
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(str(img_path), width=Inches(5))
        except Exception:
            logging.exception("Failed to embed image %s in DOCX", img_path)

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        image_match = IMAGE_PATTERN.search(stripped)
        if image_match:
            before = stripped[: image_match.start()].strip()
            if before:
                doc.add_paragraph(before)
            _add_image(image_match.group(1))
            after = stripped[image_match.end() :].strip()
            if after:
                doc.add_paragraph(after)
            continue

        if stripped.startswith("#"):
            level = 0
            for ch in stripped:
                if ch == "#":
                    level += 1
                else:
                    break
            text = stripped[level:].strip()
            level = max(1, min(level, 6))
            doc.add_heading(text or " ", level=level)
            continue

        if stripped.startswith(('- ', '* ')):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(stripped)

    docx_path = markdown_path.with_suffix(".docx")
    try:
        doc.save(str(docx_path))
        return docx_path
    except Exception:
        logging.exception("Failed to write DOCX to %s", docx_path)
        return None


MERMAID_BLOCK_RE = re.compile(r"```mermaid\s+([\s\S]*?)```", re.IGNORECASE)


def _replace_mermaid_with_images(markdown: str, markdown_path: Path) -> Tuple[str, List[Path]]:
    images: List[Path] = []
    if "```mermaid" not in markdown:
        return markdown, images
    image_dir = markdown_path.parent / "images"
    image_dir.mkdir(parents=True, exist_ok=True)

    def _render_diagram(code: str, index: int) -> Optional[Path]:
        if requests is None:
            logging.warning("requests not available; skipping mermaid rendering")
            return None
        code = code.strip()
        if not code:
            return None
        target = image_dir / f"diagram_{index}.png"
        try:
            resp = requests.post(
                "https://kroki.io/mermaid/png",
                json={"diagram_source": code},
                timeout=30,
            )
            resp.raise_for_status()
            target.write_bytes(resp.content)
            return target
        except Exception:
            logging.exception("Failed to render mermaid diagram %s", index)
            return None

    def _replace(match: re.Match) -> str:
        code = match.group(1)
        idx = len(images) + 1
        rendered = _render_diagram(code, idx)
        if rendered:
            images.append(rendered)
            rel_path = Path("images") / rendered.name
            return f"![Diagram {idx}]({rel_path.as_posix()})"
        return match.group(0)

    new_markdown = MERMAID_BLOCK_RE.sub(_replace, markdown)
    return new_markdown, images


def _export_alternate_formats(markdown_path: Path) -> Dict[str, Path]:
    outputs: Dict[str, Path] = {}
    pandoc = shutil.which("pandoc")
    if not pandoc:
        logging.warning("Pandoc not found; skipping PDF/DOCX export")
        return outputs

    targets = {
        "pdf": markdown_path.with_suffix(".pdf"),
        "docx": markdown_path.with_suffix(".docx"),
    }

    for fmt, target in targets.items():
        cmd = [pandoc, str(markdown_path), "-o", str(target)]
        try:
            subprocess.run(
                cmd,
                check=True,
                capture_output=True,
                cwd=str(markdown_path.parent),
            )
            outputs[fmt] = target
        except Exception as exc:
            logging.exception("Pandoc export failed for %s: %s", fmt, exc)
    return outputs
