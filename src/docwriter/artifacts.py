from __future__ import annotations

import base64
import io
import logging
import re
from typing import Dict, List, Optional, Tuple

try:
    import requests  # type: ignore
except Exception:  # pragma: no cover
    requests = None  # type: ignore

from .storage import BlobStore, JobStoragePaths

MERMAID_BLOCK_RE = re.compile(r"```mermaid\s+([\s\S]*?)```", re.IGNORECASE)
IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
SLUG_INVALID_RE = re.compile(r"[^a-z0-9\- ]+", re.IGNORECASE)
SLUG_SEP_RE = re.compile(r"[\s\-]+")
INTERNAL_HREF_RE = re.compile(r'href="#([A-Za-z0-9_\-]+)"')
ID_ATTR_RE = re.compile(r'id="([A-Za-z0-9_\-]+)"')
NAME_ATTR_RE = re.compile(r'name="([A-Za-z0-9_\-]+)"')


def replace_mermaid_with_images(
    markdown: str,
    job_paths: JobStoragePaths,
    store: BlobStore,
) -> Tuple[str, Dict[str, bytes]]:
    if "```mermaid" not in markdown:
        return markdown, {}

    images: Dict[str, bytes] = {}

    def _render_diagram(code: str) -> Optional[bytes]:
        if requests is None:
            logging.warning("requests not available; skipping mermaid rendering")
            return None
        code = code.strip()
        if not code:
            return None
        try:
            resp = requests.post(
                "https://kroki.io/mermaid/png",
                json={"diagram_source": code},
                timeout=30,
            )
            resp.raise_for_status()
            return resp.content
        except Exception:
            logging.exception("Failed to render mermaid diagram")
            return None

    def _replace(match: re.Match[str]) -> str:
        code = match.group(1)
        rendered = _render_diagram(code)
        if not rendered:
            return match.group(0)
        idx = len(images) + 1
        rel_path = f"images/diagram_{idx}.png"
        blob_path = job_paths.images(f"diagram_{idx}.png")
        try:
            store.put_bytes(blob=blob_path, data_bytes=rendered)
            images[rel_path] = rendered
        except Exception:
            logging.exception("Failed to upload mermaid diagram %s for job %s", idx, job_paths.job_id)
            return match.group(0)
        return f"![Diagram {idx}]({rel_path})"

    new_markdown = MERMAID_BLOCK_RE.sub(_replace, markdown)
    return new_markdown, images


def export_pdf(
    markdown: str,
    image_map: Dict[str, bytes],
    store: BlobStore,
    job_paths: JobStoragePaths,
) -> Optional[bytes]:
    try:
        html = _markdown_to_html(markdown)
        if not html:
            return None
        from weasyprint import HTML  # type: ignore
    except Exception:
        logging.warning("WeasyPrint or markdown-it-py not available; skipping PDF export")
        return None

    def _replace_src(match: re.Match[str]) -> str:
        src = match.group(1)
        data = _resolve_image_bytes(src, image_map, store, job_paths)
        if data is None:
            return match.group(0)
        encoded = base64.b64encode(data).decode("ascii")
        return f'src="data:image/png;base64,{encoded}"'

    html = re.sub(r'src="([^"]+)"', _replace_src, html)
    html = _wrap_html_for_pdf(html)
    html = _ensure_internal_anchors(html)
    try:
        return HTML(string=html).write_pdf()
    except Exception:
        logging.exception("Failed to render PDF for job %s", job_paths.job_id)
        return None


def export_docx(
    markdown: str,
    image_map: Dict[str, bytes],
    store: BlobStore,
    job_paths: JobStoragePaths,
) -> Optional[bytes]:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception:
        logging.warning("python-docx not installed; skipping DOCX export")
        return None

    doc = Document()

    def _add_image(image_src: str) -> None:
        data = _resolve_image_bytes(image_src, image_map, store, job_paths)
        if data is None:
            logging.warning("Image %s not found for DOCX export", image_src)
            return
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(data), width=Inches(5))
        except Exception:
            logging.exception("Failed to embed image %s in DOCX", image_src)

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        image_match = IMAGE_PATTERN.search(stripped)
        if image_match:
            before = stripped[: image_match.start()].strip()
            if before:
                doc.add_paragraph(before)
            _add_image(image_match.group(1))
            after = stripped[image_match.end() :].strip()
            if after:
                doc.add_paragraph(after)
            continue

        if stripped.startswith("#"):
            level = 0
            for ch in stripped:
                if ch == "#":
                    level += 1
                else:
                    break
            text = stripped[level:].strip()
            level = max(1, min(level, 6))
            doc.add_heading(text or " ", level=level)
            continue

        if stripped.startswith(('- ', '* ')):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        doc.add_paragraph(stripped)

    try:
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception:
        logging.exception("Failed to write DOCX for job %s", job_paths.job_id)
        return None


def _slugify_heading(text: str) -> str:
    normalized = SLUG_INVALID_RE.sub("", text or "")
    normalized = normalized.lower()
    parts = [part for part in SLUG_SEP_RE.split(normalized) if part]
    return "-".join(parts) or "section"


def _heading_id_plugin(md: "MarkdownIt") -> None:  # type: ignore[name-defined]
    def add_ids(state: "StateInline") -> None:  # type: ignore[name-defined]
        slug_counts: Dict[str, int] = {}
        tokens = state.tokens
        total = len(tokens)
        for idx, token in enumerate(tokens):
            if token.type != "heading_open":
                continue
            if token.attrGet("id"):
                continue
            if idx + 1 >= total:
                continue
            inline = tokens[idx + 1]
            if inline.type != "inline":
                continue
            content = inline.content.strip()
            if not content:
                continue
            base = _slugify_heading(content)
            count = slug_counts.get(base, 0)
            slug_counts[base] = count + 1
            slug = base if count == 0 else f"{base}-{count}"
            token.attrSet("id", slug)

    md.core.ruler.after("inline", "heading_ids", add_ids)


def _markdown_to_html(markdown: str) -> str:
    try:
        from markdown_it import MarkdownIt
        from markdown_it.rules_inline import StateInline
    except Exception:
        logging.warning("markdown-it-py not installed; skipping HTML conversion")
        return ""

    md = MarkdownIt("commonmark", {"html": True, "linkify": True})
    md.enable("table")
    md.enable("strikethrough")
    _heading_id_plugin(md)
    return md.render(markdown)


def _wrap_html_for_pdf(html: str) -> str:
    styles = """
    <style>
        @page {
            margin: 1in;
        }

        body {
            font-family: "Helvetica", Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.4;
        }

        img {
            max-width: 100%;
            height: auto;
            page-break-inside: avoid;
        }

        table {
            width: 100%;
            max-width: 100%;
            border-collapse: collapse;
            margin: 0.5rem 0 1rem 0;
            font-size: 12pt;
            table-layout: fixed;
        }

        th, td {
            border: 1px solid #cbd5e1;
            padding: 6px 10px;
            text-align: left;
            word-break: break-word;
            overflow-wrap: anywhere;
            white-space: normal;
        }

        th {
            background: #e2e8f0;
            font-weight: 700;
            border-top: 1px solid #cbd5e1;
            border-bottom: 2px solid #475569;
        }

        tbody tr:nth-child(odd) {
            background: #f8fafc;
        }

        tbody tr:nth-child(even) {
            background: #f1f5f9;
        }
    </style>
    """
    return f"<!DOCTYPE html><html><head><meta charset='utf-8'/>{styles}</head><body>{html}</body></html>"


def _ensure_internal_anchors(html: str) -> str:
    """Add placeholder anchors for internal href targets that are missing IDs to avoid WeasyPrint errors."""
    refs = set(INTERNAL_HREF_RE.findall(html))
    if not refs:
        return html
    existing = set(ID_ATTR_RE.findall(html)) | set(NAME_ATTR_RE.findall(html))
    missing = [ref for ref in refs if ref not in existing]
    if not missing:
        return html
    placeholders = "".join(f'<div id="{ref}" aria-hidden="true"></div>' for ref in missing)
    if "</body>" in html:
        return html.replace("</body>", f"{placeholders}</body>", 1)
    return f"{html}{placeholders}"


def _resolve_image_bytes(
    src: str,
    image_map: Dict[str, bytes],
    store: BlobStore,
    job_paths: JobStoragePaths,
) -> Optional[bytes]:
    normalized = src.lstrip("./")
    data = image_map.get(normalized)
    if data is not None:
        return data
    if normalized.startswith("http://") or normalized.startswith("https://"):
        return None
    blob_path = normalized if normalized.startswith("jobs/") else job_paths.relative(normalized)
    try:
        return store.get_bytes(blob_path)
    except Exception:
        return None
